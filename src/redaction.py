"""Dry-run and atomic production DOCX redaction workflows."""

from __future__ import annotations

import os
import re
import secrets
import tempfile
from collections import Counter
from hashlib import sha256
from pathlib import Path
from urllib.parse import unquote
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.document import Document as DocumentObject
from lxml import etree  # type: ignore[import-untyped]

from src.detector import DetectionEngine
from src.docx_audit import audit_shape_metadata
from src.docx_processor import (
    EXTRACTOR_SCHEMA_VERSION,
    TextContainer,
    iter_text_containers,
)
from src.identity_linker import IdentityLinker
from src.leak_scanner import LeakScanner
from src.models import PIIType
from src.pseudonymizer import DeterministicPseudonymizer, normalize_entity_text
from src.recognizers import all_recognizers
from src.redaction_plan import (
    PlannedReplacement,
    RedactionPlan,
    ReplacementPlanner,
    apply_redaction_plan,
)


class RedactionWriteError(ValueError):
    """Raised when a document cannot be safely written and validated."""


def _file_sha256(path: Path) -> str:
    digest = sha256()
    with path.open("rb") as source:
        for chunk in iter(lambda: source.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _validated_paths(
    input_path: Path | str, output_path: Path | str
) -> tuple[Path, Path]:
    source = Path(input_path)
    output = Path(output_path)
    if not source.is_file():
        raise FileNotFoundError(f"Input DOCX does not exist: {source}")
    if source.suffix.casefold() != ".docx":
        raise RedactionWriteError("Input path must be a .docx file")
    if output.suffix.casefold() != ".docx":
        raise RedactionWriteError("Output path must be a .docx file")
    if source.resolve() == output.resolve():
        raise RedactionWriteError("Output path must not overwrite the input DOCX")
    if output.exists():
        if output.is_dir():
            raise RedactionWriteError("Output path refers to a directory")
        if source.samefile(output):
            raise RedactionWriteError("Output path must not overwrite the input DOCX")
    return source, output


def _expected_affected_texts(
    containers: tuple[TextContainer, ...],
    plan: RedactionPlan,
) -> dict[str, str]:
    originals = {container.id: container.text for container in containers}
    grouped: dict[str, list[PlannedReplacement]] = {}
    for replacement in plan.replacements:
        grouped.setdefault(replacement.container_id, []).append(replacement)

    expected: dict[str, str] = {}
    for container_id, replacements in grouped.items():
        value = originals[container_id]
        for item in sorted(
            replacements,
            key=lambda candidate: (candidate.entity.start, candidate.entity.end),
            reverse=True,
        ):
            if value[item.entity.start : item.entity.end] != item.entity.text:
                raise RedactionWriteError(
                    "Replacement plan no longer matches its source container"
                )
            value = (
                value[: item.entity.start] + item.replacement + value[item.entity.end :]
            )
        expected[container_id] = value
    return expected


def _normalized_phone_uri(value: str) -> str:
    prefix = "+" if value.strip().startswith("+") else ""
    return prefix + re.sub(r"\D", "", value)


def _rewrite_external_pii_relationships(
    document: DocumentObject,
    replacements: tuple[PlannedReplacement, ...],
) -> int:
    """Rewrite exact mailto/tel relationships corresponding to planned PII."""

    email_replacements: dict[str, str] = {}
    phone_replacements: dict[str, str] = {}
    for item in replacements:
        if item.entity.entity_type is PIIType.EMAIL:
            email_replacements[
                normalize_entity_text(PIIType.EMAIL, item.entity.text)
            ] = item.replacement
        elif item.entity.entity_type is PIIType.PHONE:
            phone_replacements[
                normalize_entity_text(PIIType.PHONE, item.entity.text)
            ] = item.replacement

    update_count = 0
    for part in document.part.package.parts:
        for relationship in part.rels.values():
            if not relationship.is_external:
                continue
            target = relationship.target_ref
            lower_target = target.casefold()
            replacement_target: str | None = None
            if lower_target.startswith("mailto:"):
                address_and_query = target[len("mailto:") :]
                address, separator, query = address_and_query.partition("?")
                normalized = normalize_entity_text(PIIType.EMAIL, unquote(address))
                replacement = email_replacements.get(normalized)
                if replacement is not None:
                    replacement_target = f"mailto:{replacement}"
                    if separator:
                        replacement_target += f"?{query}"
            elif lower_target.startswith("tel:"):
                number = unquote(target[len("tel:") :])
                normalized = normalize_entity_text(PIIType.PHONE, number)
                replacement = phone_replacements.get(normalized)
                if replacement is not None:
                    replacement_target = f"tel:{_normalized_phone_uri(replacement)}"

            if replacement_target is not None and replacement_target != target:
                relationship._target = replacement_target  # type: ignore[attr-defined]
                update_count += 1
    return update_count


def _validate_docx_package(path: Path) -> dict[str, int | bool]:
    """Validate ZIP CRCs, required parts, and every XML relationship part."""

    try:
        with ZipFile(path) as package:
            corrupt_member = package.testzip()
            if corrupt_member is not None:
                raise RedactionWriteError("Output DOCX contains a corrupt ZIP member")
            names = set(package.namelist())
            required = {"[Content_Types].xml", "word/document.xml"}
            if not required <= names:
                raise RedactionWriteError(
                    "Output DOCX is missing required package parts"
                )
            xml_parts = [
                name
                for name in names
                if name.endswith(".xml") or name.endswith(".rels")
            ]
            for part_name in xml_parts:
                etree.fromstring(package.read(part_name))
    except (BadZipFile, etree.XMLSyntaxError) as error:
        raise RedactionWriteError("Output DOCX package validation failed") from error
    return {
        "zip_integrity_valid": True,
        "required_parts_present": True,
        "xml_parts_validated": len(xml_parts),
    }


def _unsupported_content_audit(path: Path) -> dict[str, object]:
    with ZipFile(path) as package:
        names = package.namelist()
    raster_suffixes = {".bmp", ".gif", ".jpeg", ".jpg", ".png", ".tif", ".tiff"}
    raster_images = [
        name
        for name in names
        if name.startswith("word/media/")
        and Path(name).suffix.casefold() in raster_suffixes
    ]
    unsupported_story_parts = [
        name
        for name in names
        if name in {"word/comments.xml", "word/endnotes.xml", "word/footnotes.xml"}
    ]
    embedded_objects = [name for name in names if name.startswith("word/embeddings/")]
    return {
        "raster_image_count": len(raster_images),
        "raster_image_ocr_performed": False,
        "unsupported_story_part_count": len(unsupported_story_parts),
        "embedded_object_count": len(embedded_objects),
        "limitation": (
            "Raster-image text is not OCRed. Shape descriptions/titles and "
            "Selection Pane names are audited but not automatically rewritten. "
            "Unsupported story parts and embedded objects, if counted, require "
            "manual review."
        ),
    }


def _scan_external_relationships(
    document: DocumentObject,
    *,
    detector: DetectionEngine,
    replacements: tuple[PlannedReplacement, ...],
    minimum_confidence: float = 0.85,
) -> dict[str, object]:
    synthetic_keys = {
        (
            item.entity.entity_type,
            normalize_entity_text(item.entity.entity_type, item.replacement),
        )
        for item in replacements
    }
    all_counts: Counter[str] = Counter()
    unclassified_counts: Counter[str] = Counter()
    relationship_count = 0
    high_confidence_count = 0
    known_synthetic_count = 0
    for part in document.part.package.parts:
        for relationship in part.rels.values():
            if not relationship.is_external:
                continue
            relationship_count += 1
            target = unquote(relationship.target_ref)
            for entity in detector.detect(target):
                if entity.confidence < minimum_confidence:
                    continue
                high_confidence_count += 1
                all_counts[entity.entity_type.value] += 1
                key = (
                    entity.entity_type,
                    normalize_entity_text(entity.entity_type, entity.text),
                )
                if key in synthetic_keys:
                    known_synthetic_count += 1
                else:
                    unclassified_counts[entity.entity_type.value] += 1
    return {
        "privacy_safe": True,
        "external_relationship_count": relationship_count,
        "high_confidence_detection_count": high_confidence_count,
        "known_synthetic_detection_count": known_synthetic_count,
        "unclassified_residual_count": sum(unclassified_counts.values()),
        "detections_by_type": dict(sorted(all_counts.items())),
        "unclassified_by_type": dict(sorted(unclassified_counts.items())),
    }


def build_dry_run_report(
    input_path: Path | str,
    *,
    secret: bytes | None = None,
    key_source: str = "ephemeral_dry_run",
) -> dict[str, object]:
    """Build a complete plan without mutating or saving the source document."""

    source = Path(input_path)
    if not source.is_file():
        raise FileNotFoundError(f"Input DOCX does not exist: {source}")
    before_hash = _file_sha256(source)
    effective_secret = secret or secrets.token_bytes(32)
    document = Document(str(source))
    containers = tuple(iter_text_containers(document))
    planner = ReplacementPlanner(
        DetectionEngine(all_recognizers()),
        DeterministicPseudonymizer(effective_secret),
        identity_linker=IdentityLinker(effective_secret),
    )
    plan = planner.build(containers)
    after_hash = _file_sha256(source)
    return {
        "report_schema_version": "1.0",
        "source_name": source.name,
        "source_sha256": before_hash,
        "source_unchanged": before_hash == after_hash,
        "extractor_schema_version": EXTRACTOR_SCHEMA_VERSION,
        "pseudonym_key_source": key_source,
        **plan.summary(),
    }


def redact_docx(
    input_path: Path | str,
    output_path: Path | str,
    *,
    secret: bytes,
    key_source: str = "provided",
) -> dict[str, object]:
    """Plan, mutate, atomically save, reopen, and audit a redacted DOCX."""

    source, output = _validated_paths(input_path, output_path)
    source_hash_before = _file_sha256(source)
    document = Document(str(source))
    containers = tuple(iter_text_containers(document))
    original_container_ids = {container.id for container in containers}
    original_run_count = sum(
        len(container.runs) + sum(len(mirror.runs) for mirror in container.mirrors)
        for container in containers
    )
    detector = DetectionEngine(all_recognizers())
    plan = ReplacementPlanner(
        detector,
        DeterministicPseudonymizer(secret),
        identity_linker=IdentityLinker(secret),
    ).build(containers)
    if plan.conflicts:
        reasons = Counter(item.reason for item in plan.conflicts)
        raise RedactionWriteError(
            "Redaction aborted before mutation: "
            f"{len(plan.conflicts)} structural conflict(s) "
            f"({dict(sorted(reasons.items()))})"
        )

    expected_texts = _expected_affected_texts(containers, plan)
    apply_redaction_plan(containers, plan)
    mutated_by_id = {container.id: container for container in containers}
    if any(
        mutated_by_id[item].text != expected
        for item, expected in expected_texts.items()
    ):
        raise RedactionWriteError("In-memory mutated text failed plan verification")
    relationship_update_count = _rewrite_external_pii_relationships(
        document, plan.replacements
    )
    if _file_sha256(source) != source_hash_before:
        raise RedactionWriteError("Source DOCX changed during in-memory redaction")

    output.parent.mkdir(parents=True, exist_ok=True)
    file_descriptor, temporary_name = tempfile.mkstemp(
        dir=output.parent,
        prefix=f".{output.name}.",
        suffix=".tmp.docx",
    )
    os.close(file_descriptor)
    temporary_path = Path(temporary_name)
    try:
        document.save(str(temporary_path))
        package_validation = _validate_docx_package(temporary_path)
        reopened_document = Document(str(temporary_path))
        reopened_containers = tuple(iter_text_containers(reopened_document))
        reopened_by_id = {container.id: container for container in reopened_containers}
        if set(reopened_by_id) != original_container_ids:
            raise RedactionWriteError("Container identities changed after save/reopen")
        if any(
            reopened_by_id[item].text != expected
            for item, expected in expected_texts.items()
        ):
            raise RedactionWriteError("Saved output text failed plan verification")
        reopened_run_count = sum(
            len(container.runs) + sum(len(mirror.runs) for mirror in container.mirrors)
            for container in reopened_containers
        )
        if reopened_run_count != original_run_count:
            raise RedactionWriteError("Word run count changed unexpectedly after save")

        leak_report = LeakScanner(detector).scan_containers(
            reopened_containers,
            planned_replacements=plan.replacements,
        )
        relationship_scan = _scan_external_relationships(
            reopened_document,
            detector=detector,
            replacements=plan.replacements,
        )
        unsupported_content = _unsupported_content_audit(temporary_path)
        shape_audit = audit_shape_metadata(
            temporary_path,
            detector=DetectionEngine(all_recognizers()),
        ).to_dict()
        source_hash_after = _file_sha256(source)
        if source_hash_after != source_hash_before:
            raise RedactionWriteError("Source DOCX changed while saving output")
        os.replace(temporary_path, output)
    finally:
        temporary_path.unlink(missing_ok=True)

    _validate_docx_package(output)
    Document(str(output))
    summary = plan.summary()
    replacements_by_story = Counter(item.story_type.value for item in plan.replacements)
    summary.update(
        {
            "dry_run": False,
            "document_mutated": True,
            "document_write_enabled": True,
            "quality_gate": (
                "WRITE_COMPLETED_WITH_RESIDUAL_REVIEW"
                if leak_report.review_required_count
                or relationship_scan["unclassified_residual_count"]
                else "WRITE_COMPLETED_SUPPORTED_TEXT_SCAN_CLEAR"
            ),
        }
    )
    return {
        "report_schema_version": "1.0",
        "privacy_safe": True,
        "source_name": source.name,
        "source_sha256": source_hash_before,
        "source_unchanged": _file_sha256(source) == source_hash_before,
        "output_path": str(output),
        "output_sha256": _file_sha256(output),
        "output_size_bytes": output.stat().st_size,
        "extractor_schema_version": EXTRACTOR_SCHEMA_VERSION,
        "pseudonym_key_source": key_source,
        **summary,
        "applied_replacement_count": len(plan.replacements),
        "external_relationship_update_count": relationship_update_count,
        "package_validation": package_validation,
        "post_save_verification": {
            "reopened_successfully": True,
            "container_id_set_preserved": True,
            "all_affected_container_text_verified": True,
            "run_count_preserved": True,
            "applied_replacements_by_story": dict(
                sorted(replacements_by_story.items())
            ),
        },
        "residual_supported_text_scan": leak_report.to_dict(),
        "residual_external_relationship_scan": relationship_scan,
        "unsupported_content": unsupported_content,
        "shape_metadata_audit": shape_audit,
    }


def secret_from_environment(
    variable_name: str,
    *,
    required: bool = False,
) -> tuple[bytes | None, str]:
    value = os.environ.get(variable_name)
    if value is None:
        if required:
            raise ValueError(
                f"{variable_name} is required for a write-enabled redaction"
            )
        return None, "ephemeral_dry_run"
    secret = value.encode()
    if len(secret) < 16:
        raise ValueError(f"{variable_name} must contain at least 16 UTF-8 bytes")
    return secret, f"environment:{variable_name}"
