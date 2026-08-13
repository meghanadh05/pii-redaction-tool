from __future__ import annotations

from hashlib import sha256
from pathlib import Path

import pytest
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from lxml import etree  # type: ignore[import-untyped]

from src.docx_processor import StoryKind, iter_text_containers
from src.detector import DetectionEngine
from src.leak_scanner import LeakScanner
from src.main import main
from src.models import PIIEntity, PIIType
from src.recognizers.base import Recognizer
from src.redaction import RedactionWriteError, redact_docx, secret_from_environment
from src.redaction_plan import PlannedReplacement


SECRET = b"production-write-path-test-key"
MC_NAMESPACE = "http://schemas.openxmlformats.org/markup-compatibility/2006"


class PartialSyntheticRecognizer(Recognizer):
    name = "partial_synthetic_test"
    supported_types = frozenset({PIIType.PERSON})

    def detect(self, text: str) -> list[PIIEntity]:
        value = "Person"
        start = text.find(value)
        if start < 0:
            return []
        return [
            PIIEntity(
                PIIType.PERSON,
                value,
                start,
                start + len(value),
                0.99,
                self.name,
            )
        ]


def _textbox_content(run_texts: list[str]) -> object:
    content = OxmlElement("w:txbxContent")
    paragraph = OxmlElement("w:p")
    for value in run_texts:
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = value
        run.append(text)
        paragraph.append(run)
    content.append(paragraph)
    return content


def _add_alternate_textbox(paragraph: Paragraph, value: str) -> None:
    alternate = etree.Element(
        f"{{{MC_NAMESPACE}}}AlternateContent",
        nsmap={"mc": MC_NAMESPACE},
    )
    choice = etree.SubElement(alternate, f"{{{MC_NAMESPACE}}}Choice")
    choice.set("Requires", "wps")
    choice.append(_textbox_content([value[:5], value[5:]]))
    fallback = etree.SubElement(alternate, f"{{{MC_NAMESPACE}}}Fallback")
    fallback.append(_textbox_content([value]))
    paragraph._p.append(alternate)


def test_write_path_redacts_hyperlink_reopens_and_preserves_source(
    tmp_path: Path,
) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "redacted.docx"
    original_email = "old.account@example.test"
    document = Document()
    paragraph = document.add_paragraph("Contact ")
    relationship_id = paragraph.part.relate_to(
        f"mailto:{original_email}",
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    first_run = OxmlElement("w:r")
    first_text = OxmlElement("w:t")
    first_text.text = "old.account@"
    first_run.append(first_text)
    second_run = OxmlElement("w:r")
    second_text = OxmlElement("w:t")
    second_text.text = "example.test"
    second_run.append(second_text)
    hyperlink.extend((first_run, second_run))
    paragraph._p.append(hyperlink)
    document.save(str(source))
    source_hash = sha256(source.read_bytes()).hexdigest()

    report = redact_docx(source, output, secret=SECRET, key_source="test")

    assert sha256(source.read_bytes()).hexdigest() == source_hash
    assert output.is_file()
    assert report["source_unchanged"] is True
    assert report["applied_replacement_count"] == 1
    assert report["external_relationship_update_count"] == 1
    assert report["cross_run_replacement_count"] == 1
    assert report["package_validation"]["zip_integrity_valid"] is True  # type: ignore[index]
    assert (
        report["residual_supported_text_scan"][  # type: ignore[index]
            "unclassified_residual_count"
        ]
        == 0
    )

    reopened = Document(str(output))
    containers = tuple(iter_text_containers(reopened))
    assert original_email not in "\n".join(item.text for item in containers)
    target = reopened.paragraphs[0].part.rels[relationship_id].target_ref
    assert target.startswith("mailto:synthetic.contact+")
    assert target.endswith("@example.com")
    assert original_email not in target


def test_write_path_updates_paired_textbox_representations(tmp_path: Path) -> None:
    source = tmp_path / "textbox-source.docx"
    output = tmp_path / "textbox-redacted.docx"
    original_email = "box.owner@example.test"
    document = Document()
    _add_alternate_textbox(document.add_paragraph("Outer"), original_email)
    document.save(str(source))

    report = redact_docx(source, output, secret=SECRET, key_source="test")

    assert report["text_box_replacement_count"] == 1
    assert report["mirrored_representation_update_count"] == 1
    textboxes = [
        container
        for container in iter_text_containers(Document(str(output)))
        if container.story_type is StoryKind.TEXT_BOX_PARAGRAPH
    ]
    assert len(textboxes) == 1
    assert textboxes[0].text != original_email
    assert textboxes[0].mirrors[0].logical.text == textboxes[0].text


def test_write_path_refuses_input_overwrite_and_requires_key(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "source.docx"
    Document().save(str(source))

    with pytest.raises(RedactionWriteError, match="must not overwrite"):
        redact_docx(source, source, secret=SECRET)

    monkeypatch.delenv("MISSING_REDACTION_KEY", raising=False)
    with pytest.raises(ValueError, match="required"):
        secret_from_environment("MISSING_REDACTION_KEY", required=True)


def test_write_path_fails_before_output_on_unsafe_run_content(tmp_path: Path) -> None:
    source = tmp_path / "unsafe.docx"
    output = tmp_path / "must-not-exist.docx"
    document = Document()
    run = document.add_paragraph().add_run("unsafe.owner@example.test")
    instruction = OxmlElement("w:instrText")
    instruction.text = "UNSUPPORTED"
    run._r.append(instruction)
    document.save(str(source))

    with pytest.raises(RedactionWriteError, match="structural conflict"):
        redact_docx(source, output, secret=SECRET)

    assert not output.exists()


def test_leak_scan_classifies_partial_detection_inside_synthetic_span() -> None:
    document = Document()
    paragraph = document.add_paragraph("Synthetic Person")
    container = next(iter(iter_text_containers(document)))
    original = "Alice Smith"
    entity = PIIEntity(
        PIIType.PERSON,
        original,
        0,
        len(original),
        0.95,
        "test_source",
    )
    planned = PlannedReplacement(
        container_id=container.id,
        story_type=StoryKind.BODY_PARAGRAPH,
        entity=entity,
        replacement=paragraph.text,
        affected_run_count=1,
        mirror_count=0,
    )

    report = LeakScanner(
        DetectionEngine([PartialSyntheticRecognizer()])
    ).scan_containers((container,), planned_replacements=(planned,))

    assert report.synthetic_contained_count == 1
    assert report.unclassified_residual_count == 0


def test_cli_write_mode_accepts_output_and_saves_docx(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    capsys: pytest.CaptureFixture[str],
) -> None:
    source = tmp_path / "cli-source.docx"
    output = tmp_path / "cli-redacted.docx"
    document = Document()
    document.add_paragraph("Contact cli.owner@example.test")
    document.save(str(source))
    monkeypatch.setenv("TEST_REDACTION_KEY", SECRET.decode())

    exit_code = main(
        [
            "redact",
            str(source),
            "--output",
            str(output),
            "--key-env",
            "TEST_REDACTION_KEY",
        ]
    )

    assert exit_code == 0
    assert output.is_file()
    assert '"document_write_enabled": true' in capsys.readouterr().out
