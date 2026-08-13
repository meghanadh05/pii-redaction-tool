from __future__ import annotations

from hashlib import sha256
from pathlib import Path
from zipfile import ZipFile

from docx import Document

from src.docx_processor import iter_text_containers


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SOURCE = PROJECT_ROOT / "input" / "Red Herring Prospectus.docx"
OUTPUT = PROJECT_ROOT / "output" / "Red_Herring_Prospectus_Redacted.docx"


def test_assignment_output_is_valid_distinct_docx_with_preserved_structure() -> None:
    assert OUTPUT.is_file()
    assert sha256(SOURCE.read_bytes()).hexdigest() == (
        "8b5c93f7642d659e64b51be9f6172c86c2825417f376ca1800ed331515e6f929"
    )
    assert (
        sha256(OUTPUT.read_bytes()).hexdigest()
        != sha256(SOURCE.read_bytes()).hexdigest()
    )
    with ZipFile(OUTPUT) as package:
        assert package.testzip() is None
        assert "word/document.xml" in package.namelist()

    source_document = Document(str(SOURCE))
    output_document = Document(str(OUTPUT))
    source_containers = tuple(iter_text_containers(source_document))
    output_containers = tuple(iter_text_containers(output_document))

    assert len(source_document.sections) == len(output_document.sections) == 85
    assert len(source_document.tables) == len(output_document.tables) == 76
    assert [item.id for item in source_containers] == [
        item.id for item in output_containers
    ]
    assert sum(len(item.runs) for item in source_containers) == sum(
        len(item.runs) for item in output_containers
    )
