"""Submission audit against non-prospectus synthetic DOCX documents."""

from __future__ import annotations

import json
import os
import subprocess
import sys
from collections import Counter
from pathlib import Path
from typing import Any, Callable

from docx import Document
from docx.document import Document as DocumentObject

from src.detector import DetectionEngine
from src.docx_processor import StoryKind, iter_text_containers
from src.models import PIIType
from src.pseudonymizer import replacement_passes_validator
from src.recognizers import all_recognizers


PROJECT_ROOT = Path(__file__).resolve().parents[1]
TEST_KEY_ENV = "SYNTHETIC_GENERALIZATION_KEY"
TEST_KEY = "synthetic-generalization-key-2026"
STRUCTURED_TYPES = {
    PIIType.EMAIL,
    PIIType.PHONE,
    PIIType.SSN,
    PIIType.CREDIT_CARD,
    PIIType.DOB,
    PIIType.IP_ADDRESS,
}


def _support_ticket() -> tuple[DocumentObject, tuple[str, ...], tuple[str, ...]]:
    document = Document()
    document.add_heading("SUPPORT REQUEST", level=1)
    document.add_paragraph("Ticket ID: TKT-2026-004281")
    document.add_paragraph("Order Number: ORD-771240")
    document.add_paragraph("Opened: December 10, 2025")

    for _ in range(2):
        paragraph = document.add_paragraph("Contact Person: ")
        paragraph.add_run("ANAYA ").italic = True
        paragraph.add_run("VARMAN").italic = True

    for label in ("Contact email: ", "Escalation email: "):
        paragraph = document.add_paragraph(label)
        paragraph.add_run("support.agent@").bold = True
        paragraph.add_run("example.test").bold = True
    document.add_paragraph("Phone: +1 202-555-0147")
    return (
        document,
        ("ANAYA VARMAN", "support.agent@example.test", "+1 202-555-0147"),
        (
            "TKT-2026-004281",
            "ORD-771240",
            "December 10, 2025",
            "SUPPORT REQUEST",
        ),
    )


def _employee_record() -> tuple[DocumentObject, tuple[str, ...], tuple[str, ...]]:
    document = Document()
    document.add_heading("EMPLOYEE RECORD", level=1)
    table = document.add_table(rows=6, cols=2)
    rows = (
        ("Employee", "Contact Person: DEVIKA SENVAR"),
        ("SSN", "SSN: 123-45-6789"),
        ("Birth", "DOB: 12 May 1988"),
        (
            "Mailing",
            "Mailing Address: Unit 42, Example Road, Bengaluru 560001, India",
        ),
        ("Network", "Workstation IP: 203.0.113.17"),
        ("Employee ID", "EMP-001-8872"),
    )
    for row, values in zip(table.rows, rows, strict=True):
        for cell, value in zip(row.cells, values, strict=True):
            cell.text = value
    document.add_paragraph("Review date: December 10, 2025")
    return (
        document,
        (
            "DEVIKA SENVAR",
            "123-45-6789",
            "12 May 1988",
            "Unit 42, Example Road, Bengaluru 560001, India",
            "203.0.113.17",
        ),
        ("EMP-001-8872", "December 10, 2025", "EMPLOYEE RECORD"),
    )


def _business_letter() -> tuple[DocumentObject, tuple[str, ...], tuple[str, ...]]:
    document = Document()
    section = document.sections[0]
    section.header.paragraphs[0].text = "Example Meridian Services Private Limited"
    section.footer.paragraphs[0].text = "letter.office@example.test"
    document.add_heading("GENERAL TERMS AND CONDITIONS", level=1)
    document.add_paragraph("Company: Example Meridian Services Private Limited")
    document.add_paragraph(
        "Registered Office: Flat 12, Synthetic Avenue, Pune 411001, India"
    )
    document.add_paragraph("Contact Person: MEERA KAPURAN")
    document.add_paragraph("Invoice Number: INV-880731")
    document.add_paragraph("Reference Number: REF-55-2049")
    document.add_paragraph("Offer amount: ₹4,200.00 million")
    document.add_paragraph("Page 250")
    document.add_paragraph("Securities Act provisions apply.")
    return (
        document,
        (
            "Example Meridian Services Private Limited",
            "letter.office@example.test",
            "Flat 12, Synthetic Avenue, Pune 411001, India",
            "MEERA KAPURAN",
        ),
        (
            "INV-880731",
            "REF-55-2049",
            "₹4,200.00 million",
            "Page 250",
            "GENERAL TERMS AND CONDITIONS",
            "Securities Act provisions apply.",
        ),
    )


def _international_contacts() -> tuple[
    DocumentObject, tuple[str, ...], tuple[str, ...]
]:
    document = Document()
    document.add_heading("INTERNATIONAL CONTACTS", level=1)
    document.add_paragraph("Company: Example Northstar Trading Ltd.")
    document.add_paragraph("Contact Person: KAVYA SENVAR")
    document.add_paragraph("Email: global.desk@example.test")
    document.add_paragraph("Phone: +44 7700 900123")
    document.add_paragraph(
        "Office Address: Unit 7, Example Street, Northstar Building, "
        "London, United Kingdom"
    )
    document.add_paragraph("Payment card: 4242 4242 4242 4242")
    document.add_paragraph("IPv6: 2001:db8:ffff::7")
    document.add_paragraph("Reference Number: INT-2026-88115")
    document.add_paragraph("Document date: December 10, 2025")
    return (
        document,
        (
            "Example Northstar Trading Ltd.",
            "KAVYA SENVAR",
            "global.desk@example.test",
            "+44 7700 900123",
            ("Unit 7, Example Street, Northstar Building, London, United Kingdom"),
            "4242 4242 4242 4242",
            "2001:db8:ffff::7",
        ),
        ("INT-2026-88115", "December 10, 2025", "INTERNATIONAL CONTACTS"),
    )


def _run_public_cli(source: Path, output: Path) -> dict[str, Any]:
    environment = os.environ.copy()
    environment[TEST_KEY_ENV] = TEST_KEY
    result = subprocess.run(
        [
            sys.executable,
            "-m",
            "src.main",
            "redact",
            str(source),
            "--output",
            str(output),
            "--key-env",
            TEST_KEY_ENV,
        ],
        cwd=PROJECT_ROOT,
        env=environment,
        check=True,
        capture_output=True,
        text=True,
        timeout=90,
    )
    return json.loads(result.stdout)


def _all_story_text(document: DocumentObject) -> str:
    return "\n".join(container.text for container in iter_text_containers(document))


def test_public_cli_generalizes_across_four_synthetic_docx_fixtures(
    tmp_path: Path,
) -> None:
    builders: tuple[
        tuple[
            str, Callable[[], tuple[DocumentObject, tuple[str, ...], tuple[str, ...]]]
        ],
        ...,
    ] = (
        ("synthetic_ticket", _support_ticket),
        ("synthetic_employee", _employee_record),
        ("synthetic_letter", _business_letter),
        ("synthetic_international", _international_contacts),
    )
    aggregate_planned: Counter[str] = Counter()
    aggregate_output_entities: Counter[PIIType] = Counter()
    reports: list[dict[str, Any]] = []

    for stem, builder in builders:
        document, source_pii, preserved_negatives = builder()
        source = tmp_path / f"{stem}.docx"
        output = tmp_path / f"{stem}_redacted.docx"
        document.save(str(source))

        report = _run_public_cli(source, output)
        reports.append(report)
        assert output.is_file()
        assert report["source_unchanged"] is True
        assert report["package_validation"]["zip_integrity_valid"] is True
        assert report["post_save_verification"]["reopened_successfully"] is True
        assert report["conflict_count"] == 0
        assert (
            report["residual_supported_text_scan"]["unclassified_residual_count"] == 0
        )
        aggregate_planned.update(report["planned_replacements_by_type"])

        reopened = Document(str(output))
        output_text = _all_story_text(reopened)
        for value in source_pii:
            assert value not in output_text
        for value in preserved_negatives:
            assert value in output_text

        entities = [
            entity
            for container in iter_text_containers(reopened)
            for entity in DetectionEngine(all_recognizers()).detect(container.text)
        ]
        aggregate_output_entities.update(entity.entity_type for entity in entities)
        for entity in entities:
            if entity.entity_type in STRUCTURED_TYPES:
                assert replacement_passes_validator(entity.entity_type, entity.text)

    assert all(aggregate_planned[entity_type.value] > 0 for entity_type in PIIType)
    assert all(aggregate_output_entities[entity_type] > 0 for entity_type in PIIType)
    assert sum(report["cross_run_replacement_count"] for report in reports) >= 2
    assert any(
        report["post_save_verification"]["applied_replacements_by_story"].get(
            StoryKind.TABLE_CELL_PARAGRAPH.value, 0
        )
        > 0
        for report in reports
    )
    assert any(
        report["post_save_verification"]["applied_replacements_by_story"].get(
            StoryKind.HEADER_PARAGRAPH.value, 0
        )
        > 0
        for report in reports
    )
    assert any(
        report["post_save_verification"]["applied_replacements_by_story"].get(
            StoryKind.FOOTER_PARAGRAPH.value, 0
        )
        > 0
        for report in reports
    )

    ticket_output = Document(str(tmp_path / "synthetic_ticket_redacted.docx"))
    ticket_paragraphs = [paragraph.text for paragraph in ticket_output.paragraphs]
    repeated_people = [
        value.split(": ", 1)[1]
        for value in ticket_paragraphs
        if value.startswith("Contact Person: ")
    ]
    repeated_emails = [
        value.split(": ", 1)[1]
        for value in ticket_paragraphs
        if value.startswith(("Contact email: ", "Escalation email: "))
    ]
    assert len(set(repeated_people)) == 1
    assert len(set(repeated_emails)) == 1

    formatted_email_paragraph = next(
        paragraph
        for paragraph in ticket_output.paragraphs
        if paragraph.text.startswith("Contact email: ")
    )
    synthetic_email_run = next(
        run for run in formatted_email_paragraph.runs if "@example.com" in run.text
    )
    assert synthetic_email_run.bold is True
