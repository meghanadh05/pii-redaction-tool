from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from src.detector import DetectionEngine
from src.docx_processor import StoryKind, TextContainer
from src.models import PIIEntity, PIIType
from src.pseudonymizer import DeterministicPseudonymizer
from src.recognizers.base import Recognizer
from src.redaction import build_dry_run_report
from src.redaction_plan import ReplacementPlanner, apply_redaction_plan


SECRET = b"replacement-planning-test-key"


class LiteralRecognizer(Recognizer):
    name = "literal_test"
    supported_types = frozenset({PIIType.PERSON, PIIType.EMAIL, PIIType.COMPANY})
    values = {
        "Alice Smith": PIIType.PERSON,
        "alice.smith@real.test": PIIType.EMAIL,
        "Real Company Limited": PIIType.COMPANY,
    }

    def detect(self, text: str) -> list[PIIEntity]:
        return [
            PIIEntity(entity_type, value, match.start(), match.end(), 0.99, self.name)
            for value, entity_type in self.values.items()
            for match in re.finditer(re.escape(value), text)
        ]


class OverlappingDetector:
    def detect(self, text: str) -> list[PIIEntity]:
        return [
            PIIEntity(PIIType.PERSON, text[0:11], 0, 11, 0.9, "overlap"),
            PIIEntity(PIIType.COMPANY, text[6:15], 6, 15, 0.9, "overlap"),
        ]


class LocallyQualifiedPersonRecognizer(Recognizer):
    name = "locally_qualified_person_test"
    supported_types = frozenset({PIIType.PERSON})

    def detect(self, text: str) -> list[PIIEntity]:
        value = "Alice Smith"
        if not text.startswith("Contact Person:"):
            return []
        start = text.index(value)
        return [
            PIIEntity(
                PIIType.PERSON,
                value,
                start,
                start + len(value),
                0.95,
                self.name,
            )
        ]


def container(
    run_texts: list[str],
    container_id: str,
    story: StoryKind,
    *,
    mirror: bool = False,
) -> tuple[TextContainer, Paragraph | None]:
    document = Document()
    paragraph = document.add_paragraph()
    for text in run_texts:
        paragraph.add_run(text)
    mirror_paragraph = None
    if mirror:
        mirror_paragraph = document.add_paragraph("".join(run_texts))
    return (
        TextContainer.from_paragraph(
            paragraph,
            container_id=container_id,
            story_type=story,
            mirror_paragraphs=(mirror_paragraph,) if mirror_paragraph else (),
        ),
        mirror_paragraph,
    )


def test_plan_is_non_mutating_consistent_and_applies_all_story_types() -> None:
    body, _ = container(["Alice ", "Smith"], "body/p0000", StoryKind.BODY_PARAGRAPH)
    table, _ = container(
        ["Alice Smith"],
        "body/t0000/r0000/c0000/p0000",
        StoryKind.TABLE_CELL_PARAGRAPH,
    )
    header, _ = container(
        ["alice.smith@real.test"],
        "header:/word/header1.xml:default/p0000",
        StoryKind.HEADER_PARAGRAPH,
    )
    textbox, mirror = container(
        ["Real ", "Company Limited"],
        "body/p0001/txb0000/p0000",
        StoryKind.TEXT_BOX_PARAGRAPH,
        mirror=True,
    )
    containers = (body, table, header, textbox)
    originals = tuple(item.text for item in containers)
    plan = ReplacementPlanner(
        DetectionEngine([LiteralRecognizer()]),
        DeterministicPseudonymizer(SECRET),
    ).build(containers)

    assert tuple(item.text for item in containers) == originals
    assert plan.conflicts == ()
    assert len(plan.replacements) == 4
    names = [
        item.replacement
        for item in plan.replacements
        if item.entity.entity_type is PIIType.PERSON
    ]
    assert len(set(names)) == 1
    assert plan.summary()["cross_run_replacement_count"] == 2
    assert plan.summary()["mirrored_representation_update_count"] == 1
    safe_json = json.dumps(plan.summary())
    assert "Alice Smith" not in safe_json
    assert "Real Company Limited" not in safe_json

    apply_redaction_plan(containers, plan)

    assert all("Alice Smith" not in item.text for item in (body, table))
    assert textbox.text == plan.replacements[-1].replacement
    assert mirror is not None
    assert mirror.text == textbox.text


def test_overlapping_plan_conflict_is_detected_before_mutation() -> None:
    target, _ = container(["Alice Smith Ltd"], "body/p0000", StoryKind.BODY_PARAGRAPH)
    original = target.text
    planner = ReplacementPlanner(
        OverlappingDetector(),  # type: ignore[arg-type]
        DeterministicPseudonymizer(SECRET),
    )

    plan = planner.build((target,))

    assert len(plan.conflicts) == 1
    assert plan.conflicts[0].reason == "OVERLAPPING_DETECTIONS"
    assert plan.replacements == ()
    assert target.text == original


def test_small_docx_dry_run_does_not_write_or_expose_values(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("Contact alice.smith@real.test")
    document.save(str(source))
    before = source.read_bytes()

    report = build_dry_run_report(source, secret=SECRET, key_source="test")

    assert source.read_bytes() == before
    assert report["source_unchanged"] is True
    assert report["document_mutated"] is False
    assert report["planned_replacements_by_type"]["EMAIL"] == 1  # type: ignore[index]
    serialized = json.dumps(report)
    assert "alice.smith@real.test" not in serialized


def test_plan_propagates_high_confidence_exact_repeats_without_overlap() -> None:
    labelled, _ = container(
        ["Contact Person: Alice Smith; minutes signed by Alice Smith"],
        "body/p0000",
        StoryKind.BODY_PARAGRAPH,
    )
    unrelated, _ = container(
        ["A separate narrative mentions Alice Smith"],
        "body/p0001",
        StoryKind.BODY_PARAGRAPH,
    )
    containers = (labelled, unrelated)
    plan = ReplacementPlanner(
        DetectionEngine([LocallyQualifiedPersonRecognizer()]),
        DeterministicPseudonymizer(SECRET),
    ).build(containers)

    assert len(plan.replacements) == 2
    assert plan.summary()["exact_repeat_propagation_count"] == 1
    assert len({item.replacement for item in plan.replacements}) == 1

    apply_redaction_plan(containers, plan)

    assert "Alice Smith" not in labelled.text
    assert "Alice Smith" in unrelated.text
