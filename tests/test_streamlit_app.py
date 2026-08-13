"""Focused privacy and interaction checks for the Streamlit frontend."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any, cast

from docx import Document
from streamlit.testing.v1 import AppTest

from src.models import PIIType
from streamlit_app import load_blind_evaluation, mask_entity


PROJECT_ROOT = Path(__file__).resolve().parents[1]


def _synthetic_docx() -> bytes:
    document = Document()
    paragraph = document.add_paragraph("Contact email: ")
    paragraph.add_run("test.owner@").bold = True
    paragraph.add_run("example.test").bold = True
    content = BytesIO()
    document.save(content)
    return content.getvalue()


def test_mask_entity_never_returns_complete_original() -> None:
    samples = {
        PIIType.PERSON: "Anaya Varman",
        PIIType.EMAIL: "anaya.varman@example.test",
        PIIType.PHONE: "+44 7700 900123",
        PIIType.COMPANY: "Example Meridian Services Private Limited",
        PIIType.ADDRESS: "Unit 42, Example Road, Bengaluru 560001, India",
        PIIType.SSN: "123-45-6789",
        PIIType.CREDIT_CARD: "4242 4242 4242 4242",
        PIIType.DOB: "12 May 1988",
        PIIType.IP_ADDRESS: "2001:db8:ffff::7",
    }

    for entity_type, original in samples.items():
        masked = mask_entity(entity_type, original)
        assert masked != original
        assert "•" in masked


def test_frontend_metrics_come_from_frozen_blind_artifact() -> None:
    report = load_blind_evaluation()
    micro = report["metrics"]["exact"]["micro"]

    assert report["dataset_role"] == "blind_final_evaluation_single_run"
    assert report["sample"]["reviewed_container_count"] == 180
    assert micro["tp"] == 141
    assert micro["fp"] == 68
    assert micro["fn"] == 27
    assert round(micro["precision"], 3) == 0.675
    assert round(micro["recall"], 3) == 0.839


def test_streamlit_upload_analyze_redact_verify_download_flow() -> None:
    app = AppTest.from_file(PROJECT_ROOT / "streamlit_app.py", default_timeout=90).run()
    assert not app.exception
    assert [tab.label for tab in app.tabs] == [
        "Overview",
        "Analyze",
        "Redact",
        "Verify",
        "Evaluation",
    ]

    cast(Any, app.get("file_uploader")[0]).upload(
        "synthetic-contact.docx",
        _synthetic_docx(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ).run()
    assert not app.exception
    app.button[0].click().run()

    assert not app.exception
    assert "Show original PII (unsafe)" in [toggle.label for toggle in app.toggle]
    assert "Redact and verify DOCX" in [button.label for button in app.button]
    app.button[1].click().run()

    assert not app.exception
    assert not app.error
    assert [cast(Any, button).label for button in app.get("download_button")] == [
        "Download redacted DOCX"
    ]
    verification = {
        metric.label: metric.value
        for metric in app.metric
        if metric.label in {"Replacements", "Cross-run", "Conflicts", "Residual review"}
    }
    assert verification == {
        "Replacements": "1",
        "Cross-run": "1",
        "Conflicts": "0",
        "Residual review": "0",
    }
