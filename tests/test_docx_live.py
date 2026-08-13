from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.enum.section import WD_SECTION
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from lxml import etree
from lxml.etree import _Element

from src.docx_processor import (
    EXTRACTOR_SCHEMA_VERSION,
    StoryKind,
    TextContainer,
    TextReplacement,
    UnsupportedRunContentError,
    iter_text_containers,
)


MC_NAMESPACE = "http://schemas.openxmlformats.org/markup-compatibility/2006"


def _textbox_content(run_texts: list[str]) -> _Element:
    content = OxmlElement("w:txbxContent")
    paragraph = OxmlElement("w:p")
    for value in run_texts:
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = value
        run.append(text)
        paragraph.append(run)
    content.append(paragraph)
    return content


def add_alternate_textbox(
    paragraph: Paragraph,
    *,
    choice_runs: list[str],
    fallback_runs: list[str],
) -> None:
    alternate = etree.Element(
        f"{{{MC_NAMESPACE}}}AlternateContent",
        nsmap={"mc": MC_NAMESPACE},
    )
    choice = etree.SubElement(alternate, f"{{{MC_NAMESPACE}}}Choice")
    choice.set("Requires", "wps")
    choice.append(_textbox_content(choice_runs))
    fallback = etree.SubElement(alternate, f"{{{MC_NAMESPACE}}}Fallback")
    fallback.append(_textbox_content(fallback_runs))
    paragraph._p.append(alternate)


def make_container(run_texts: list[str]) -> TextContainer:
    document = Document()
    paragraph = document.add_paragraph()
    for value in run_texts:
        paragraph.add_run(value)
    return TextContainer.from_paragraph(
        paragraph,
        container_id="body/p0000",
        story_type=StoryKind.BODY_PARAGRAPH,
    )


@pytest.mark.parametrize(
    ("run_texts", "target", "replacement", "expected_runs"),
    [
        (
            ["Email old@example.test now"],
            "old@example.test",
            "new@example.com",
            ["Email new@example.com now"],
        ),
        (
            ["Call +91 ", "9876543210 today"],
            "+91 9876543210",
            "+91 9000000000",
            ["Call +91 9000000000", " today"],
        ),
        (
            ["Dear Al", "ic", "e Smith!"],
            "Alice Smith",
            "Priya Shah",
            ["Dear Priya Shah", "", "!"],
        ),
        (["prefix-Ali", "ce-suffix"], "Alice", "Priya", ["prefix-Priya", "-suffix"]),
        (["“Al", "ice”—contact"], "Alice", "Priya", ["“Priya", "”—contact"]),
    ],
)
def test_live_rewrite_preserves_runs_prefix_suffix_and_unicode(
    run_texts: list[str],
    target: str,
    replacement: str,
    expected_runs: list[str],
) -> None:
    container = make_container(run_texts)
    start = container.text.index(target)

    container.rewrite([TextReplacement(start, start + len(target), replacement)])

    assert [run.text for run in container.runs] == expected_runs
    assert container.text == "".join(expected_runs)


def test_live_rewrite_preserves_first_affected_run_formatting() -> None:
    container = make_container(["prefix ", "Alice", " suffix"])
    container.runs[1].bold = True
    container.runs[2].italic = True
    start = container.text.index("Alice")

    container.rewrite([TextReplacement(start, start + 5, "Priya Shah")])

    assert container.runs[1].text == "Priya Shah"
    assert container.runs[1].bold is True
    assert container.runs[2].text == " suffix"
    assert container.runs[2].italic is True


def test_live_rewrite_handles_multiple_replacements_safely() -> None:
    container = make_container(["Alice: ", "old@example.test", "; Alice"])
    second_name_start = container.text.rindex("Alice")

    container.rewrite(
        [
            TextReplacement(0, 5, "Priya"),
            TextReplacement(7, 23, "new@example.com"),
            TextReplacement(second_name_start, second_name_start + 5, "Priya"),
        ]
    )

    assert container.text == "Priya: new@example.com; Priya"
    assert len(container.runs) == 3


def test_extraction_distinguishes_stories_and_deduplicates_linked_header(
    tmp_path: Path,
) -> None:
    source = tmp_path / "stories.docx"
    document = Document()
    document.add_paragraph("Body before")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Cell text"
    document.add_paragraph("Body after")
    document.sections[0].header.paragraphs[0].text = "Shared header"
    document.sections[0].footer.paragraphs[0].text = "Shared footer"
    second_section = document.add_section(WD_SECTION.NEW_PAGE)
    assert second_section.header.is_linked_to_previous is True
    assert second_section.footer.is_linked_to_previous is True
    document.save(str(source))

    reopened = Document(str(source))
    containers = tuple(iter_text_containers(reopened))
    nonempty = [container for container in containers if container.text]

    assert [container.story_type for container in nonempty] == [
        StoryKind.BODY_PARAGRAPH,
        StoryKind.TABLE_CELL_PARAGRAPH,
        StoryKind.BODY_PARAGRAPH,
        StoryKind.HEADER_PARAGRAPH,
        StoryKind.FOOTER_PARAGRAPH,
    ]
    assert len({container.id for container in containers}) == len(containers)
    assert sum(container.text == "Shared header" for container in containers) == 1
    assert sum(container.text == "Shared footer" for container in containers) == 1

    second_pass_ids = [
        container.id for container in iter_text_containers(Document(str(source)))
    ]
    assert second_pass_ids == [container.id for container in containers]


def test_table_header_and_footer_rewrites_survive_save_and_reopen(
    tmp_path: Path,
) -> None:
    output = tmp_path / "rewritten.docx"
    document = Document()
    table = document.add_table(rows=1, cols=1)
    cell_paragraph = table.cell(0, 0).paragraphs[0]
    cell_paragraph.add_run("old@").bold = True
    cell_paragraph.add_run("example.test")
    header_paragraph = document.sections[0].header.paragraphs[0]
    header_paragraph.add_run("Call +91 ")
    header_paragraph.add_run("9876543210")
    footer_paragraph = document.sections[0].footer.paragraphs[0]
    footer_paragraph.add_run("old@").italic = True
    footer_paragraph.add_run("example.test")

    containers = tuple(iter_text_containers(document))
    cell_container = next(
        container for container in containers if "old@" in container.text
    )
    header_container = next(
        container for container in containers if container.text.startswith("Call")
    )
    footer_container = next(
        container
        for container in containers
        if container.story_type is StoryKind.FOOTER_PARAGRAPH
        and "old@" in container.text
    )
    cell_container.rewrite(
        [TextReplacement(0, len(cell_container.text), "new@example.com")]
    )
    phone_start = header_container.text.index("+91")
    header_container.rewrite(
        [
            TextReplacement(
                phone_start,
                len(header_container.text),
                "+91 9000000000",
            )
        ]
    )
    footer_container.rewrite(
        [TextReplacement(0, len(footer_container.text), "new@example.com")]
    )
    document.save(str(output))

    reopened = tuple(iter_text_containers(Document(str(output))))
    assert any(container.text == "new@example.com" for container in reopened)
    assert any(container.text == "Call +91 9000000000" for container in reopened)
    assert sum(container.text == "new@example.com" for container in reopened) == 2


def test_nested_table_paragraphs_are_extracted_once() -> None:
    document = Document()
    outer = document.add_table(rows=1, cols=1)
    nested = outer.cell(0, 0).add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "Nested cell"

    containers = tuple(iter_text_containers(document))
    nested_containers = [item for item in containers if item.text == "Nested cell"]

    assert len(nested_containers) == 1
    assert "/t0000/" in nested_containers[0].id
    assert nested_containers[0].story_type is StoryKind.TABLE_CELL_PARAGRAPH


def test_unsupported_run_content_fails_before_mutation() -> None:
    container = make_container(["prefix Alice suffix"])
    instruction = OxmlElement("w:instrText")
    instruction.text = "UNSUPPORTED"
    container.runs[0]._r.append(instruction)
    original = container.runs[0].text
    start = container.text.index("Alice")

    with pytest.raises(UnsupportedRunContentError):
        container.rewrite([TextReplacement(start, start + 5, "Priya")])

    assert container.runs[0].text == original


def test_hyperlink_runs_are_extracted_and_rewritten_without_losing_relationship(
    tmp_path: Path,
) -> None:
    output = tmp_path / "hyperlink.docx"
    document = Document()
    paragraph = document.add_paragraph("Email ")
    relationship_id = paragraph.part.relate_to(
        "mailto:old@example.test",
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run_element = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = "old@example.test"
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)

    container = TextContainer.from_paragraph(
        paragraph,
        container_id="body/p0000",
        story_type=StoryKind.BODY_PARAGRAPH,
    )
    start = container.text.index("old@example.test")
    container.rewrite(
        [TextReplacement(start, start + len("old@example.test"), "new@example.com")]
    )
    document.save(str(output))

    reopened = Document(str(output))
    extracted = tuple(iter_text_containers(reopened))
    assert extracted[0].text == "Email new@example.com"
    hyperlink_elements = reopened.paragraphs[0]._p.xpath("./w:hyperlink")
    assert len(hyperlink_elements) == 1
    assert hyperlink_elements[0].get(qn("r:id")) == relationship_id


def test_choice_fallback_textbox_is_detected_once_and_rewritten_twice(
    tmp_path: Path,
) -> None:
    output = tmp_path / "textbox.docx"
    document = Document()
    outer = document.add_paragraph("Outer text")
    add_alternate_textbox(
        outer,
        choice_runs=["old@", "example.test"],
        fallback_runs=["old@example", ".test"],
    )

    containers = tuple(iter_text_containers(document))
    textboxes = [
        container
        for container in containers
        if container.story_type is StoryKind.TEXT_BOX_PARAGRAPH
    ]

    assert EXTRACTOR_SCHEMA_VERSION == "1.0"
    assert len(textboxes) == 1
    assert textboxes[0].text == "old@example.test"
    assert len(textboxes[0].mirrors) == 1
    assert textboxes[0].metadata["mirror_status"] == "choice_fallback_paired"

    textboxes[0].rewrite(
        [TextReplacement(0, len(textboxes[0].text), "new@example.com")]
    )
    assert textboxes[0].text == "new@example.com"
    assert textboxes[0].mirrors[0].logical.text == "new@example.com"
    document.save(str(output))

    reopened = tuple(iter_text_containers(Document(str(output))))
    reopened_textboxes = [
        container
        for container in reopened
        if container.story_type is StoryKind.TEXT_BOX_PARAGRAPH
    ]
    assert len(reopened_textboxes) == 1
    assert reopened_textboxes[0].text == "new@example.com"
    assert reopened_textboxes[0].mirrors[0].logical.text == "new@example.com"


def test_mismatched_choice_fallback_is_visible_but_not_rewritable() -> None:
    document = Document()
    outer = document.add_paragraph()
    add_alternate_textbox(
        outer,
        choice_runs=["Choice value"],
        fallback_runs=["Fallback value"],
    )

    textboxes = [
        container
        for container in iter_text_containers(document)
        if container.story_type is StoryKind.TEXT_BOX_PARAGRAPH
    ]

    assert [container.text for container in textboxes] == [
        "Choice value",
        "Fallback value",
    ]
    assert all(container.metadata["rewrite_safe"] == "false" for container in textboxes)
    with pytest.raises(UnsupportedRunContentError):
        textboxes[0].rewrite([TextReplacement(0, len(textboxes[0].text), "New")])
